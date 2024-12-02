#!/usr/bin/env python3

__author__ = "ChatGPT"


import os
from docx import Document
from docx.shared import Inches

def create_final_report(results_files, roc_curves, output_file):
    # initialization
    doc = Document()
    doc.add_heading("Final Model Comparison Report", level=1)

    doc.add_paragraph(f"A total of {len(results_files)} models were compared.")

    model_data = {}
    # Find results from files
    for result_file in results_files:
        model_name = os.path.basename(result_file).replace("_results.txt", "")
        with open(result_file, "r") as f:
            lines = f.readlines()

        accuracy_lines = [line for line in lines if "accuracy:" in line.lower()]
        if not accuracy_lines:
            print(f"Warning: No 'accuracy:' found in {result_file}. Skipping.")
            continue

        try:
            accuracy = float(accuracy_lines[0].split(":")[1].strip())
        except (IndexError, ValueError) as e:
            print(f"Error parsing accuracy in {result_file}: {e}")
            continue

        # save models data
        model_data[model_name] = {
            "accuracy": accuracy,
            "description": "".join(lines),
            "roc_curve": next((curve for curve in roc_curves if model_name in curve), None),
        }

    # sorting
    sorted_models = sorted(model_data.items(), key=lambda x: x[1]["accuracy"], reverse=True)

    # add table with accuracy
    doc.add_heading("Model Accuracies", level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid"
    table.rows[0].cells[0].text = "Model"
    table.rows[0].cells[1].text = "Accuracy"
    for model_name, data in sorted_models:
        row = table.add_row().cells
        row[0].text = model_name
        row[1].text = f"{data['accuracy']:.4f}"

    # select best model
    best_model, best_data = sorted_models[0]
    doc.add_heading(f"Best Model: {best_model}", level=2)
    doc.add_paragraph(best_data["description"])

    # add roc-curve
    if best_data["roc_curve"] and os.path.exists(best_data["roc_curve"]):
        doc.add_heading("ROC Curve", level=3)
        doc.add_picture(best_data["roc_curve"], width=Inches(5))

    # save doc file
    doc.save(output_file)

create_final_report(snakemake.input.results, snakemake.input.roc_curves, snakemake.output[0])
