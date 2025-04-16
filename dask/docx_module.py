#!/usr/bin/env python3

'''
These module provide functionality for creating a Word document report.
The report includes performance analysis of Pandas and Dask.
Generated with ChatGPT and modified by V.Pazenko.
'''

__author__ = "V.Pazenko"
__version__ = 1.0

from docx import Document
from docx.shared import Inches
import re
import time
import text_module


def parse_log(log_file, start_time):
    '''
    Input: 1. log_file = path to the log file
           2. start_time = time starting from which data from the log_file should be processed

    Function parses the log_file and extracts relevant information for Pandas and Dask.

    Output: two dictionaries: one for Pandas and one for Dask, containing the relevant metrics.
    '''
    pandas_data = {}
    dask_data = {}

    with open(log_file, "r") as f:
        for line in f:
            match = re.match(r"(\d{4}-\d{2}-\d{2} \d{2}:\d{2}:\d{2}),\d+ - (.*)", line)
            if match:
                timestamp, message = match.groups()
                log_time = time.mktime(time.strptime(timestamp, "%Y-%m-%d %H:%M:%S"))
                if log_time < start_time:
                    continue

                if "Pandas" in message:
                    key, value = message.replace("Pandas ", "").split(": ")
                    pandas_data[key.split(". ")[1]] = value.strip()
                elif "Dask" in message:
                    key, value = message.replace("Dask ", "").split(": ")
                    dask_data[key.split(". ")[1]] = value.strip()

    return pandas_data, dask_data


def create_docx_report(log_file, start_time, grouped_df, docx_name="Report.docx"):
    '''
    Input: 1. log_file = path to the log file
           2. start_time = time starting from which data from the log_file should be processed
           3. grouped_df = pd.DataFrame with the average expression of selected genes
           4. docx_name = name of the output Word document

    Function create docx file and fill it with the data from the log_file.
    The report includes performance analysis of Pandas and Dask.

    Output: creates a Word document with the performance analysis.
    '''
    # Parse log file
    pandas_data, dask_data = parse_log(log_file, start_time)
    # Create a new Word document
    doc = Document()
    doc.add_heading('Performance Analysis', level=1)

    doc.add_paragraph("This document provides an analysis of performance metrics for Pandas and Dask.")

    # Time table
    doc.add_heading('Execution Time Comparison', level=2)
    table = doc.add_table(rows=7, cols=3)
    table.style = 'Table Grid'

    headers = ['Process', 'Pandas Time (s)', 'Dask Time (s)']
    for i, header in enumerate(headers):
        table.cell(0, i).text = header

    processes = [
        "load time", "preprocessing time", "data exploration time", 
        "preprocessing (train/test) time", "modelling time", "Total time (measured)"
    ]

    for i, process in enumerate(processes, 1):
        table.cell(i, 0).text = process.capitalize()
        table.cell(i, 1).text = pandas_data.get(process, "N/A")
        table.cell(i, 2).text = dask_data.get(process, "N/A")

    # Memory table
    doc.add_heading('Memory Usage', level=2)
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'

    table.cell(0, 0).text = 'Dataset'
    table.cell(0, 1).text = 'Pandas (MB)'
    table.cell(0, 2).text = 'Dask (MB)'

    name = 'clinical'
    val = f'memory usage ({name})'
    table.cell(1, 0).text = name.capitalize()
    table.cell(1, 1).text = pandas_data.get(val, "N/A")
    table.cell(1, 2).text = dask_data.get(val, "N/A")

    name = 'expression'
    val = f'memory usage for the first partition ({name})'
    table.cell(2, 0).text = ''.join([name.capitalize(), ' (partitioned)'])
    table.cell(2, 1).text = pandas_data.get(val, "N/A")
    table.cell(2, 2).text = dask_data.get(val, "N/A")

    table.cell(3, 0).text = name.capitalize()
    table.cell(3, 1).text = pandas_data.get(f"memory usage ({name})", "N/A")
    table.cell(3, 2).text = dask_data.get(f"memory usage approximately computed ({name})", "N/A")

    # add some text
    text_module.add_text_to_doc(doc, 'Execution time')

    # Models table
    doc.add_heading('Model Performance Comparison', level=2)
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'

    metrics = ['accuracy', 'precision', 'recall', 'F1-score', 'AUC-ROC']

    table.cell(0, 0).text = 'Metric'
    table.cell(0, 1).text = 'Pandas'
    table.cell(0, 2).text = 'Dask'

    for i, metric in enumerate(metrics, 1):
        table.cell(i, 0).text = metric.capitalize()
        table.cell(i, 1).text = pandas_data.get(f'Model {metric}', "N/A")
        table.cell(i, 2).text = dask_data.get(f'Model {metric}', "N/A")

    # Add some text
    text_module.add_text_to_doc(doc, 'Model')

    # Add 2 roc curves in one line
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    cell1, cell2 = table.rows[0].cells
    roc_pd = "./plots/ROC_curve_pandas.png"
    roc_dask = "./plots/ROC_curve_dask.png"
    cell1.paragraphs[0].add_run().add_picture(roc_pd, width=Inches(3))
    cell2.paragraphs[0].add_run().add_picture(roc_dask, width=Inches(3))
    text_module.add_text_to_doc(doc, 'concl')
    
    # Go to the New page
    doc.add_page_break()
    doc.add_paragraph("Group patients by clinical variables and compute the average expression of selected genes")

    # Add pic
    image_path = "./plots/TumorSubtype_distribution_pandas.png"  # Укажите путь к изображению
    doc.add_picture(image_path, width=Inches(4))

    # Add table with average expression of selected genes from pd df
    doc.add_heading('Average expression of top 10 variable genes within TumorSubtype groups', level=2)
    table.style = 'Table Grid'

    grouped_df.iloc[:, 1:] = grouped_df.iloc[:, 1:].apply(lambda x: round(x, 2))
    table = doc.add_table(rows=grouped_df.shape[0] + 1, cols=grouped_df.shape[1])
    
    for j, column in enumerate(grouped_df.columns):
        table.cell(0, j).text = column

    for i, row in grouped_df.iterrows():
        for j, value in enumerate(row):
            if j == 0:
                table.cell(i + 1, j).text = str(int(value))
            else:
                table.cell(i + 1, j).text = str(value)

    # Save final docx file
    doc.save(docx_name)
